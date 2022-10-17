from cmd import IDENTCHARS
from winsound import SND_LOOP
from intWeb import storage, login_required_auth
from intWeb import get_assest_model
from flask import flash,Blueprint, current_app, redirect, render_template, request, \
    session, url_for,send_file,Flask,send_from_directory,jsonify
import os
import zipfile
import math
from urllib.parse import quote
from datetime import datetime,date
from flask import Response
import json
import io
from openpyxl.workbook import Workbook
from openpyxl import load_workbook   
from flask_qrcode import QRcode


qrcode=QRcode(current_app)

crud = Blueprint('crud', __name__)

def upload_hw_file(file,UPLOAD_FOLDER,seat):
    if not file:
        return None
    public_url = storage.upload_hw_file(
        file,
        file.filename,
        file.content_type,
        UPLOAD_FOLDER,
        0,
        seat
    )
    current_app.logger.info(
        "Uploaded file %s as %s.", file.filename, public_url)
    return public_url


def upload_image_file(file,UPLOAD_FOLDER,pixStr=None):
    """
    Upload the user-uploaded file to Google Cloud Storage and retrieve its
    publicly-accessible URL.
    """
    if not file:
        return None
    public_url = storage.upload_file(
        file,#.read(),
        file.filename,
        file.content_type,
        UPLOAD_FOLDER,
        pixStr,
        "EsAsset",
    )
    current_app.logger.info(
        "Uploaded file %s as %s.", file.filename, public_url)
    return public_url




@crud.route("/")
@login_required_auth
def home():
    token = request.args.get('page_token', None)
    if token:
        token = token.encode('utf-8')
    books, next_page_token = get_assest_model().list(cursor=token)
    return render_template(
        "esasset/index.html",
        books=books,
        next_page_token=next_page_token)


@crud.route("/list")
@login_required_auth
def list():
    token = request.args.get('page_token', None)
    if token:
        token = token.encode('utf-8')
    books, next_page_token = get_assest_model().list_desc(cursor=token)
    return render_template(
        "esasset/list.html",
        books=books,
        next_page_token=next_page_token)

# [START list_mine]
@crud.route("/mine")
@login_required_auth
def list_mine():
    token = request.args.get('page_token', None)
    if token:
        token = token.encode('utf-8')
    books, next_page_token = get_assest_model().list_by_user(
        user_id=session['profile']['id'],
        cursor=token)
    return render_template(
        "esasset/list.html",
        books=books,
        next_page_token=next_page_token)
# [END list_mine]


@crud.route('/acc/addbatch/<cnt>', methods=['GET', 'POST'])
@login_required_auth
def acc_addbatch(cnt):
    regSDate= (request.args.get('regSDate', datetime.today().strftime( '%Y-%m-%d')))
    data={
        "acno":str(hex(int(datetime.utcnow().timestamp()))),
        "regSDate":regSDate}

    data['regSDate']=datetime.strptime(data['regSDate'], '%Y-%m-%d')
    book = get_assest_model().createAcc_Blank(data,int(cnt))
    return f"{cnt}"
# [END add]


@crud.route('/acc/JSON/update/<itemid>', methods=['GET', 'POST'])
@login_required_auth
def acc_JsonUpdate(itemid):
    data=request.get_json()
    if 'regSDate' in data:
        data['regSDate']=datetime.strptime(data['regSDate'], '%Y-%m-%d')
    book = get_assest_model().update(data, itemid)
    return jsonify( book)


@crud.route("/showacc")
@login_required_auth
def show_acc_grid():
    token = request.args.get('page_token', None)
    if token:
        token = token.encode('utf-8')
    books, next_page_token = get_assest_model().list(limit=2000,cursor=token)
    return render_template(
        "esasset/acc_grid.html",
        
        items=books,
        next_page_token=next_page_token)  


@crud.route("/showitem")
@login_required_auth
def show_item_grid():
    token = request.args.get('page_token', None)
    if token:
        token = token.encode('utf-8')
    books, next_page_token = get_assest_model().Itemlist(limit=6000,cursor=token)
    return render_template(
        "esasset/grid.html",
        book={"id":0},
        items=books,
        next_page_token=next_page_token)  


@crud.route("/QueryForm", methods=['GET', 'POST'])
@login_required_auth
def show_QueryForm():
    token = request.args.get('page_token', None)
    if token:
        token = token.encode('utf-8')

    if request.method == 'POST':
        data = request.form.to_dict(flat=True)
        if data["TableName"]=="Item":
            books = get_assest_model().Itemlist_by_FilterOption({data["FieldName"]:data["FieldValue"]})
            return render_template(
                "esasset/grid.html",
                book={"id":0},
                items=books)


        if data["TableName"]=="Acc":
            books = get_assest_model().Acclist_by_FilterOption({data["FieldName"]:data["FieldValue"]})
            return render_template(
                "esasset/acc_grid.html",
                book={"id":0},
                items=books)

    return render_template("esasset/QueryForm.html")


@crud.route("/locationitemlist/<roomid>")
def locationitemlist(roomid):
    token = request.args.get('page_token', None)
    if token:
        token = token.encode('utf-8')
    books, next_page_token = get_assest_model().locationitemlist_desc(roomid=roomid,buwei=1000000,cursor=token)
    #between
    return render_template(
        "esasset/item/list.html",
        books=books,
        next_page_token=next_page_token,roomid=roomid)

@crud.route("/snitemlist/<serialno>")
def snitemlist(serialno):
    token = request.args.get('page_token', None)
    if token:
        token = token.encode('utf-8')
    if len(serialno)<4 : 
        return render_template("esasset/item/list.html", books=[])

    books, next_page_token = get_assest_model().snitemlist_desc(sn=serialno,cursor=token)
    #between
    return render_template(
        "esasset/item/list.html",
        books=books,
        next_page_token=next_page_token)


@crud.route("/modelitemlist/<model_no>")
def modelitemlist(model_no):
    token = request.args.get('page_token', None)
    if token:
        token = token.encode('utf-8')
    if len(model_no)<4 : 
        return render_template("esasset/item/list.html", books=[])

    books, next_page_token = get_assest_model().modelitemlist_desc(model=model_no,cursor=token)
    #between
    return render_template(
        "esasset/item/list.html",
        books=books,
        next_page_token=next_page_token)


@crud.route("/categoryitemlist/<cateid>")
def categoryitemlist(cateid):
    token = request.args.get('page_token', None)
    if token:
        token = token.encode('utf-8')
    buwei=1000000
    modwei=10000000000 
    if len(cateid)>4 : buwei=1000
    books, next_page_token = get_assest_model().categoryitemlist_asc(cateid=cateid,modwei=modwei,buwei=buwei,cursor=token)
    #between
    return render_template(
        "esasset/grid.html",
        book={"id":0},items=books,
        next_page_token=next_page_token)


@crud.route("/JSON/categoryitemlist/<cateid>")
def JSONcategoryitemlist(cateid):
    token = request.args.get('page_token', None)
    if token:
        token = token.encode('utf-8')
    buwei=1000
    mask=request.args.get("mask",None)
    if mask!=None:
        buwei=int(mask)
    books, next_page_token = get_assest_model().categoryitemlist_desc(cateid=cateid,buwei=buwei,cursor=token)
    print(books)
    return jsonify(books)


def Get_FileList(crspath, filenames, prefix=None):
    path = current_app.config['HW_UPLOAD_FOLDER']
    LECTURE_FOLDER = os.path.join(path, crspath)
    if not os.path.isdir(LECTURE_FOLDER):
        os.mkdir(LECTURE_FOLDER)
    else:
        for root,dirs, files in os.walk(LECTURE_FOLDER):
            for file in files:
                if prefix==None:
                    basename, extension = file.rsplit('.', 1)
                    _file=basename.split('-_')[0]+"."+extension
                    filenames.append({"f":quote(str(file)),"n":_file})    
                elif prefix in file:
                    basename, extension = file.rsplit('.', 1)
                    _file=basename.split('-_')[0]+"."+extension
                    filenames.append({"f":quote(str(file)),"n":_file})    
    pass

#####
# [START add]


@crud.route('/<id>/item/<itemid>')
@login_required_auth
def itemview(id,itemid):
    book,acc_id= get_assest_model().readItem(itemid)
    if id!="0" : acc_id=id
    filenames=[]
    crspath=str(book["acc_acno"])
    Get_FileList(crspath,filenames,"ASS"+crspath+"ID"+str(book["id"])+"_")             
    return render_template("esasset/item/view.html",acc_id=acc_id, book=book,filenames=filenames)


@crud.route('/<id>/item/add', methods=['GET', 'POST'])
@login_required_auth
def itemadd(id):
    acc_acno= request.args.get('acno', "0")
    sess= request.args.get('sess', "0000")
    regSDate= request.args.get('regSDate', datetime.today().strftime( '%Y-%m-%d'))
    if request.method == 'POST':
        data = request.form.to_dict(flat=True)

        # If an image was uploaded, update the data to point to the new image.
        path = current_app.config['HW_UPLOAD_FOLDER']
        image_url = upload_image_file(request.files.get('image'),path,acc_acno)

        if image_url:
            data['imageUrl'] = image_url

        # If the user is logged in, associate their profile with the new book.
        if 'profile' in session:
            data['createdById'] = session['profile']['id']
        data['regSDate']=datetime.strptime(data['regSDate'], '%Y-%m-%d')

        for f_ in ['itemno',"price", "quantity", "p_amount", "amount", "fund_amount","gno","ict"]:
            if f_ in data:   
                if data[f_]=="" or data[f_]=="None" or data[f_]==None:
                    data[f_]=None
        book = get_assest_model().createItem(data)

        return redirect(url_for('.itemview', id=id,itemid=book['id']))
    book={
        "quantity":"0",
        "price":"0",
        "adjust":"0",
        "amount":"0",
        "depr_year":"0",
        "sess":sess,
        "acc_acno":acc_acno,
        "regSDate":regSDate}
    return render_template("esasset/item/form.html", action="Add", book=book)
# [END add]


@crud.route('/<id>/item/addbatch/<cnt>', methods=['GET', 'POST'])
@login_required_auth
def itemaddbatch(id,cnt):
    acc_acno= (request.args.get('acno', "0"))
    sess= (request.args.get('sess', "0000"))
    regSDate= (request.args.get('regSDate', datetime.today().strftime( '%Y-%m-%d')))
    if acc_acno=="0" :
        acc_rec = get_assest_model().read(id)
        regSDate=acc_rec["regSDate"].strftime( '%Y-%m-%d')
        acc_acno=  acc_rec["acno"]
    data={
        "itemno":"",
        "quantity":"0",
        "price":"0",
        "amount":"0",
        "sess":sess,
        "acc_acno":acc_acno,
        "regSDate":regSDate}
        
    if 'profile' in session:
        data['createdById'] = session['profile']['id']

    data['regSDate']=datetime.strptime(data['regSDate'], '%Y-%m-%d')
    if data["itemno"]==""  or data["itemno"]=="None" or data["itemno"]==None:
        data["itemno"]=None
    else:
        pass
    #for i in range(int(cnt)):
    book = get_assest_model().createItem_Blank(data,int(cnt))
    return f"{cnt}"
# [END add]


@crud.route('/<id>/item/api/JSON/update/<itemid>', methods=['GET', 'POST'])
@login_required_auth
def itemJsonUpdate(id,itemid):
    data=request.get_json()
    if 'regSDate' in data:
        data['regSDate']=datetime.strptime(data['regSDate'], '%Y-%m-%d')
    for f_ in ['itemno','itemcatno',"price", "quantity", "p_amount", "amount", "fund_amount","gno","ict"]:
        if f_ in data:   
            if data[f_]=="" or data[f_]=="None" or data[f_]==None:
                data[f_]=None
    print(data)
    print(itemid)
    book = get_assest_model().updateItem(data, itemid)
    return jsonify( book)


@crud.route('/<id>/item/api/JSON/updateSet/<nothing>', methods=['GET', 'POST'])
@login_required_auth
def itemJsonUpdateSet(id,nothing):
    data=request.get_json()
    for itemid in data:
        if 'regSDate' in data[itemid]:
            data[itemid]['regSDate']=datetime.strptime(data[itemid]['regSDate'], '%Y-%m-%d')
        for f_ in ['itemno','itemcatno',"price", "quantity", "p_amount", "amount", "fund_amount","gno","ict"]:
            
            if f_ in data[itemid]:
                if data[itemid][f_]=="" or data[itemid][f_]=="None" or data[itemid][f_]==None:
                    data[itemid][f_]=None
    cnt = get_assest_model().updateItem_DataSet(data)
    return f"updated {cnt} rows."


def CheckOwnRecordErr(book,session):
    if session["profile"].get("id")==1 : return None
    if str(session["profile"].get("id")) != book["createdById"] :  return "no create user!"
    return None


@crud.route('/<id>/item/<itemid>/edit', methods=['GET', 'POST'])
@login_required_auth
def itemedit(id,itemid):
    book, acc_id  = get_assest_model().readItem(itemid)
    Err=CheckOwnRecordErr(book,session) 
    if Err!= None :  return Err
    if  book['regSDate'] !=None:
        book['regSDate']=book['regSDate'].isoformat()[:10]
    if request.method == 'POST':
        data = request.form.to_dict(flat=True)
        path = current_app.config['HW_UPLOAD_FOLDER']
        image_url = upload_image_file(request.files.get('image'),path,str(data["acc_acno"]))
        if image_url:
            data['imageUrl'] = image_url
        
        data['regSDate']=datetime.strptime(data['regSDate'], '%Y-%m-%d')
        if data["itemno"]=="" or data["itemno"]=="None" or data["itemno"]==None:
            data["itemno"]=None
        
        for f_ in ['itemno',"itemcatno","price", "quantity", "p_amount", "amount", "fund_amount","gno","ict"]:
            if f_ in data:   
                if data[f_]=="" or data[f_]=="None" or data[f_]==None:
                    data[f_]=None      
        
        book = get_assest_model().updateItem(data, itemid)
        #return redirect(url_for('.view', id=book['id']))
        return redirect(url_for('.itemview', id=id,itemid=book['id']))

    return render_template("esasset/item/form.html", action="Edit", book=book)


#####
@crud.route('/<id>/item/<itemid>/delete')
@login_required_auth
def itemdelete(id,itemid):
    book,acc_id = get_assest_model().readItem(itemid)
    crspath=book["acc_acno"]
    if (book["createdById"]==str(session['profile']['id']))  :
        path = current_app.config['HW_UPLOAD_FOLDER']
        #UPLOAD_FOLDER = os.path.join(path, crspath)
        #for root,dirs, files in os.walk(UPLOAD_FOLDER):
        #   for file in files:      
        #       os.remove(UPLOAD_FOLDER+"/"+file)
        get_assest_model().deleteItem(itemid)        
    return redirect(url_for('.view',id=id))


@crud.route('/<id>/item/<itemid>/upload', methods=['GET', 'POST'])
def uploadItemfiles(id,itemid):
    book,acc_id = get_assest_model().readItem(itemid)
    crspath=str(book["acc_acno"])
    path = current_app.config['HW_UPLOAD_FOLDER']
    UPLOAD_FOLDER = os.path.join(path, crspath)
    if not os.path.isdir(UPLOAD_FOLDER):
        os.mkdir(UPLOAD_FOLDER)
    if request.method == 'POST':
        if 'files[]' not in request.files:
            flash('No file part')
            return render_template("view.html", book=book)
        files = request.files.getlist('files[]')
        for file in files:
            upload_hw_file(file,UPLOAD_FOLDER,"ASS"+crspath+"ID"+str(book["id"])+"_")
        flash('File(s) successfully uploaded')
        return redirect(f"/EsAsset/{id}")
        #return render_template("view.html", book=book)    


@crud.route('/<id>/item/<itemid>/download/<filename>')
def download_item_file(id,itemid,filename):
    book,acc_id = get_assest_model().readItem(itemid)
    crspath=str(book["acc_acno"])
    #return render_template("view.html", book=book)
    # Get current path os.getcwd()
    path = current_app.config['HW_UPLOAD_FOLDER']
    # file Upload
    UPLOAD_FOLDER = os.path.join(path, crspath)
    FilePath=UPLOAD_FOLDER+"/"+filename
    basename, extension = filename.rsplit('.', 1)
    _file=basename.split('-_')[0]+"."+extension    
    if extension in  ['jpg']:
        return send_file(FilePath,
                mimetype = 'image/*')
    else:
        return send_file(FilePath,
                mimetype = 'zip',
                attachment_filename= _file,
                as_attachment = True)
    # Delete the zip file if not needed


#####
@crud.route('/<id>')
@login_required_auth
def view(id):
    if id=="0" : return redirect("/EsAsset/")
    book = get_assest_model().read(id)
    netdev=book["netdev"]
    book["regSDate"]=book["regSDate"].strftime( '%Y-%m-%d')
    items=  get_assest_model().Itemlist_by_acno(book["acno"])
    filenames=[]
    crspath=str(book["acno"])
    Get_FileList(crspath,filenames)     #"ACC"+crspath+"_"
    return render_template("esasset/view.html",netdev=netdev, book=book,items=items,filenames=filenames)


# [START add]
@crud.route('/add', methods=['GET', 'POST'])
@login_required_auth
def add():
    if request.method == 'POST':
        data = request.form.to_dict(flat=True)
        # If an image was uploaded, update the data to point to the new image.
        path = current_app.config['HW_UPLOAD_FOLDER']
        image_url = upload_image_file(request.files.get('image'),path,str(data["acno"]))
        if image_url:
            data['imageUrl'] = image_url
        # If the user is logged in, associate their profile with the new book.
        if 'profile' in session:
            data['createdById'] = session['profile']['id']
        data['regSDate']=datetime.strptime(data['regSDate'], '%Y-%m-%d')
        book = get_assest_model().create(data)
        return redirect(url_for('.view', id=book['id']))
    book={
        "total":"0",
        "readonly":"0",
        "regSDate":datetime.today().strftime( '%Y-%m-%d')
        }
    return render_template("esasset/form.html", action="Add", book=book)
# [END add]


@crud.route('/<id>/edit', methods=['GET', 'POST'])
@login_required_auth
def edit(id):
    book = get_assest_model().read(id)
    Err=CheckOwnRecordErr(book,session) 
    if Err!= None :  return Err
    book["regSDate"]=book["regSDate"].strftime( '%Y-%m-%d')
    if request.method == 'POST':
        data = request.form.to_dict(flat=True)
        path = current_app.config['HW_UPLOAD_FOLDER']
        image_url = upload_image_file(request.files.get('image'),path,str(data["acno"]))
        if image_url:
            data['imageUrl'] = image_url
        data['regSDate']=datetime.strptime(data['regSDate'], '%Y-%m-%d')
        book = get_assest_model().update(data, id)
        return redirect(url_for('.view', id=book['id']))
    return render_template("esasset/form.html", action="Edit", book=book)


@crud.route('/<id>/delete')
@login_required_auth
def delete(id):
    book = get_assest_model().read(id)
    Err=CheckOwnRecordErr(book,session)
    if Err != None:  return Err
    crspath=book["Path"]
    if (book["createdById"]==str(session['profile']['id']))  :
        path = current_app.config['HW_UPLOAD_FOLDER']
        UPLOAD_FOLDER = os.path.join(path, crspath)
        for root,dirs, files in os.walk(UPLOAD_FOLDER):
           for file in files:      
               os.remove(UPLOAD_FOLDER+"/"+file)
        get_assest_model().delete(id)        
    return redirect(url_for('.list'))


@crud.route("/<id>/itemgrid")
@login_required_auth
def itemgrid(id):
    book = get_assest_model().read(id)
    Err=CheckOwnRecordErr(book,session)
    if Err != None:  return Err
    netdev=book["netdev"]
    book["regSDate"]=book["regSDate"].strftime( '%Y-%m-%d')
    items= get_assest_model().Itemlist_by_acno(book["acno"])
    filenames=[]
    return render_template("esasset/grid.html",netdev=netdev, book=book,items=items,lecturesfile=[],filenames=filenames)

@crud.route("/<id>/itemgridMTCode")
@login_required_auth
def itemgridMTCode(id):
    book = get_assest_model().read(id)
    Err=CheckOwnRecordErr(book,session)
    if Err != None:  return Err
    #netdev=book["netdev"]
    #book["regSDate"]=book["regSDate"].strftime( '%Y-%m-%d')
    items= get_assest_model().Itemlist_by_acno(book["acno"])
    #filenames=[]
    #return render_template("esasset/grid.html",netdev=netdev, book=book,items=items,lecturesfile=[],filenames=filenames)
    ###
    document = Document("doc/lebal_templ_schsub951.docx")
    template = document.tables[0]
    tbl = template._tbl
    lebals_len=len(items)
    for i in range(lebals_len):
        new_tbl = copy.deepcopy(tbl)
        # Then we do the replacement
        # replaceText(document, '<>', 'New value')
        #paragrah = document.add_paragraph()
        paragraph =document.paragraphs[0]
        # After that, we add the previously copied table
        paragraph._p.addnext(new_tbl)    

    for i, table in enumerate(document.tables):
        if i>=len(items):            
            break
        lebal=items[i]
        qr_txt=lebal["fund_lebal"]
        lebal_txt_1=f"{lebal['fund_name']}\n{lebal['name']}\n{lebal['fund_lebal']}"
        lebal_txt_2=f"{lebal['fund_name']}\n{lebal['name']}\n{lebal['sn']}\n{lebal['fund_lebal']}"
        hdr_cells = table.rows[0].cells
        file = lebalQrcode.make(qr_txt)
        buf = io.BytesIO()
        file.save(buf)
        buf.seek(0)
        p=hdr_cells[0].paragraphs[0]
        run=p.add_run()
        run.add_picture(buf, width=Cm(1.8)) # width=Inches(1.25))
        p=hdr_cells[1].paragraphs[0]
        if lebal["sn"]==None or lebal["sn"]=="":
            cell_run=p.add_run(lebal_txt_1)            
            cell_run.font.name = '微軟正黑體'
        else:
            cell_run=p.add_run(lebal_txt_2)
            cell_run.font.name = '新細明體'#'微軟正黑體'            
        cell_run.font.size = Pt(8)
        #cell_run.bold = True    
    f = io.BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    return send_file(f, as_attachment=True, attachment_filename='report.doc')        
    ###

@crud.route("/<id>/itemgridMTCodeAH")
@login_required_auth
def itemgridMTCodeAH(id):
    book = get_assest_model().read(id)
    Err=CheckOwnRecordErr(book,session)
    if Err != None:  return Err
    items= get_assest_model().Itemlist_by_acno(book["acno"])
    document = Document("doc/lebal_templ_schsub951.docx")
    template = document.tables[0]
    tbl = template._tbl
    lebals_len=len(items)
    for i in range(lebals_len):
        new_tbl = copy.deepcopy(tbl)
        # Then we do the replacement
        # replaceText(document, '<>', 'New value')
        #paragrah = document.add_paragraph()
        paragraph =document.paragraphs[0]
        # After that, we add the previously copied table
        paragraph._p.addnext(new_tbl)    

    for i, table in enumerate(document.tables):
        if i>=len(items):            
            break
        lebal=items[i]
        qr_txt=lebal["fund_lebal"]
        lebal_ctx=lebal['name']
        if lebal["sn"]==None or lebal["sn"]=="":
            pass
        else:
            lebal_ctx=lebal_ctx+'\n'+lebal['sn']
        lebal_txt=[lebal['fund_name'],lebal_ctx,lebal['fund_lebal']]
        hdr_cells = table.rows[0].cells
        file = lebalQrcode.make(qr_txt)
        buf = io.BytesIO()
        file.save(buf)
        buf.seek(0)
        p=hdr_cells[0].paragraphs[0]
        run=p.add_run()
        run.add_picture(buf, width=Cm(1.8)) # width=Inches(1.25))

        p=hdr_cells[1].paragraphs[0]
        for idx, line_ctx in enumerate(lebal_txt):
            cell_run=p.add_run(line_ctx+'\n')    
            if idx==2 or idx==2:
                cell_run.font.name = '新細明體'#'微軟正黑體'  #cell_run.font.name = '新細明體'#'微軟正黑體'            
                cell_run.font.size = Pt(8)
                cell_run.bold = True    
            else:
                cell_run.font.name = '新細明體'#'微軟正黑體'            
                cell_run.font.size = Pt(8)
                #cell_run.bold = True    

    f = io.BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    return send_file(f, as_attachment=True, attachment_filename='report.doc')        
    ###


@crud.route("/<id>/DownloadXLS",methods=["GET"])
def get_Acc_DownloadXLS(id):
    f2A=["-","itemcatno","-","-","-","sess","vouchernum","regSDate","invoicenum","name","model","sn","supplier","quantity","price","p_amount","place","-","fund_amount","fund_name","keeper","amount","depr_year","warr_period","note1","note2"]
    datestr= datetime.today().strftime( '%Y-%m-%d')
    wb = load_workbook(filename = 'doc/asset_page_templ.xlsx')
    cate_row_idx={}
    sn_="page"
    book = get_assest_model().read(id)    
    wb[sn_]["I3"]=book["acno"]
    wb[sn_]["I4"]=book["acc"]
    items= get_assest_model().Itemlist_by_acno(book["acno"])
    cate_row_idx[sn_]=7
    for r_ in items:
        ridx=cate_row_idx[sn_]
        for i, f_ in enumerate(f2A):
            if f_ == "itemcatno":
                itemcatno=str(r_[f_]).zfill(14)
                wb[sn_][f"{chr(65+i)}{ridx}"]= r_["fund_lebal"] if itemcatno== "0000000000None" else itemcatno
            elif f_ == "regSDate":
                wb[sn_][f"{chr(65+i)}{ridx}"]=r_[f_].strftime( '%Y-%m-%d')
            elif f_ != "-":
                wb[sn_][f"{chr(65+i)}{ridx}"]=r_[f_]
            if i==15 :
                wb[sn_][f"{chr(65+i)}{ridx}"].number_format = '#,##0.00'
        wb[sn_].row_dimensions[ridx].height = 30                   
        cate_row_idx[sn_]=ridx+1
    file = io.BytesIO()
    wb.save(file)
    file.seek(0)
    return send_file(file, attachment_filename=f"asset_page_{datestr}.xlsx", as_attachment=True)


@crud.route("/itemCateGrid")
@login_required_auth
def itemCateGrid():
    items= get_assest_model().ItemCatlist()
    return render_template("esasset/itemCategory/grid.html",netdev=1, items=items)
    

@crud.route('/itemCateGrid_/addbatch/<cnt>', methods=['GET', 'POST'])
@login_required_auth
def itemcataddbatch(cnt):
    items= get_assest_model().createItemCat_blank(int(cnt))
    return f"{cnt}"


@crud.route('/itemCateGrid_/api/JSON/updateSet/<nothing>', methods=['GET', 'POST'])
@login_required_auth
def itemCateGridJsonUpdateSet(nothing):
    data=request.get_json()
    book = get_assest_model().updateItemCat_DataSet(data)
    return "update !"



@crud.route('/itemCateGrid_/api/JSON/update/<itemcat_id>', methods=['GET', 'POST'])
@login_required_auth
def itemCateGridJsonUpdate(itemcat_id):
    data=request.get_json()
    book = get_assest_model().updateItemCat(data, itemcat_id)
    return jsonify( book)
    

@crud.route('/itemCateGrid_/api/OUTJSON', methods=['GET', 'POST'])
@login_required_auth
def itemCateGridApiOUTJSON():
    items= get_assest_model().ItemCatlist()
    return jsonify(items)


@crud.route('/api/JSON/pushItemnoMoveLog', methods=['GET', 'POST'])
@login_required_auth
def pushItemMoveLogJson():
    data=request.get_json()
    if 'profile' in session:
        data['createdById'] = session['profile']['id']
    print(data)
    book = get_assest_model().createItemMoveLog(data)
    return jsonify( book)


@crud.route('/api/JSON/getItemnoMoveLog_by_itemno/<itemno>', methods=['GET', 'POST'])
@login_required_auth
def getItemMoveLogJson(itemno):
    book = get_assest_model().ItemMoveLoglist_by_itemno(itemno)
    return jsonify( book)


@crud.route('/<id>/downloadall')
def download_all(id):
    book = get_assest_model().read(id)
    crspath=book["Path"]
    seat=session['profile']['Seat']
    path = current_app.config['HW_UPLOAD_FOLDER']
    UPLOAD_FOLDER = os.path.join(path, crspath)
    ZIP_PATH = os.path.join(path, "ZIPFILE")
    if not os.path.isdir(ZIP_PATH):
        os.mkdir(ZIP_PATH)
    # Zip file Initialization and you can change the compression type
    ZipFileName=f"HW{crspath}.zip"
    ZipFilePath=ZIP_PATH+"/"+ZipFileName
    zipfolder = zipfile.ZipFile(ZipFilePath,'w', compression = zipfile.ZIP_STORED)
    # zip all the files which are inside in the folder
    for root,dirs, files in os.walk(UPLOAD_FOLDER):
        for file in files:
            zipfolder.write(UPLOAD_FOLDER+'/'+file)
    zipfolder.close()
    return send_file(ZipFilePath,
            mimetype = 'zip',
            attachment_filename= ZipFileName,
            as_attachment = True)
    #os.remove(ZipFilePath)

@crud.route('/<id>/download/<filename>')
def download_file(id,filename):
    book = get_assest_model().read(id)
    crspath=str(book["acno"])
    #return render_template("view.html", book=book)
    # Get current path os.getcwd()
    path = current_app.config['HW_UPLOAD_FOLDER']
    # file Upload
    UPLOAD_FOLDER = os.path.join(path, crspath)
    FilePath=UPLOAD_FOLDER+"/"+filename
    basename, extension = filename.rsplit('.', 1)
    _file=basename.split('-_')[0]+"."+extension    
    if extension in  ['jpg']:
        return send_file(FilePath,
                mimetype = 'image/*',
                )
    else:
        return send_file(FilePath,
                mimetype = 'zip',
                attachment_filename= _file,
                as_attachment = True)
    # Delete the zip file if not needed

@crud.route('/<id>/downloadlecture/<filename>')
def download_lecturefile(id,filename):
    book = get_assest_model().read(id)
    crspath=book["Path"]
    #return render_template("view.html", book=book)
    # Get current path os.getcwd()
    path = current_app.config['HW_UPLOAD_FOLDER']
    # file Upload
    UPLOAD_FOLDER = os.path.join(path, crspath+"LECTURE")
    FilePath=UPLOAD_FOLDER+"/"+filename
    basename, extension = filename.rsplit('.', 1)
    _file=basename.split('-_')[0]+"."+extension    
    return send_file(FilePath,
            mimetype = 'zip',
            attachment_filename= _file,
            as_attachment = True)
    # Delete the zip file if not needed


@crud.route("/<id>/gridMTImg")
@login_required_auth
def gridMTImg(id):
    book = get_assest_model().read(id)
    Err=CheckOwnRecordErr(book,session)
    if Err != None:  return Err
    items= get_assest_model().Itemlist_by_acno(book["acno"])
    return render_template("esasset/gridimg.html", book=book,items=items)


@crud.route('/MTImg/<mtcode>')
def showMTImg(mtcode):
    # Get current path os.getcwd()
    path = current_app.config['HW_UPLOAD_FOLDER']
    fix=request.args.get("fix", default="")
    if fix=="" :
        pass
    else:
        fix="_"+fix
    # file Upload
    #UPLOAD_FOLDER = os.path.join(path, filename)
    FilePath=path+"/MT/"+mtcode+".jpg"
    if os.path.exists(FilePath):
        return send_file(FilePath,
            attachment_filename=f"{mtcode}{fix}.jpg",
            mimetype = 'image/*')
    else:
        return "error"
    # Delete the zip file if not needed

@crud.route('/<id>/img/<filename>')
def showimage(id,filename):
    # Get current path os.getcwd()
    path = current_app.config['HW_UPLOAD_FOLDER']
    # file Upload
    #UPLOAD_FOLDER = os.path.join(path, filename)
    FilePath=path+"/"+filename
    return send_file(FilePath,
            mimetype = 'image/*')
    # Delete the zip file if not needed

@crud.route('/<id>/upload', methods=['GET', 'POST'])
def uploadfiles(id):
    book = get_assest_model().read(id)
    crspath=str(book["acno"])
    path = current_app.config['HW_UPLOAD_FOLDER']
    UPLOAD_FOLDER = os.path.join(path, crspath)
    if not os.path.isdir(UPLOAD_FOLDER):
        os.mkdir(UPLOAD_FOLDER)
    if request.method == 'POST':
        if 'files[]' not in request.files:
            flash('No file part')
            return render_template("view.html", book=book)
        
        files = request.files.getlist('files[]')
        for file in files:
            upload_hw_file(file,UPLOAD_FOLDER,"ACC"+crspath+"_")
        
        flash('File(s) successfully uploaded')
        return redirect(f"/EsAsset/{id}")
        #return render_template("view.html", book=book)


@crud.route('/<id>/cleanclasswork')
@login_required_auth
def cleanclasswork(id):
    book = get_assest_model().read(id)
    Err=CheckOwnRecordErr(book,session)
    if Err != None:  return Err
    crspath=book["Path"]
    if (book["createdById"]==str(session['profile']['id']))  :
        path = current_app.config['HW_UPLOAD_FOLDER']
        UPLOAD_FOLDER = os.path.join(path, crspath)
        for root,dirs, files in os.walk(UPLOAD_FOLDER):
           for file in files:      
               os.remove(UPLOAD_FOLDER+"/"+file)
        #get_assest_model().delete(id)        
    return redirect(url_for('.list'))

@crud.route('/ZIPALLPIC')
@login_required_auth
def ZIPALLPIC():
    path = current_app.config['HW_UPLOAD_FOLDER']
    ZIP_PATH =current_app.config['HW_TEMP_FOLDER']
    if not os.path.isdir(ZIP_PATH):
        os.mkdir(ZIP_PATH)
    # Zip file Initialization and you can change the compression type
    ZipFileName=f"HW.zip"
    ZipFilePath=ZIP_PATH+""+ZipFileName
    zipfolder = zipfile.ZipFile(ZipFilePath,'w', compression = zipfile.ZIP_STORED)
    # zip all the files which are inside in the folder
    for root,dirs, files in os.walk(path):
        for file in files:
            file_path=root+"/"+file
            file_path=file_path.replace("//","/")
            zipfolder.write(file_path)
    zipfolder.close()
    return send_file(ZipFilePath,
            mimetype = 'zip',
            attachment_filename= ZipFileName,
            as_attachment = True)
    #os.remove(ZipFilePath)    



@crud.route('/JSON/db/<tablename>')
@login_required_auth
def JSON_DB(tablename):
    book = get_assest_model().readAllFromTable(tablename)
    return jsonify(book)


def convert_timestamp(item_date_object):
    if isinstance(item_date_object, (date, datetime)):
        return item_date_object.timestamp()

@crud.route('/JSON/file/<tablename>')
@login_required_auth
def JSON_DBFILE(tablename):
    book = get_assest_model().readAllFromTable(tablename)
    xml =json.dumps(book,default=convert_timestamp)
    return Response(xml, mimetype='text/json',content_type="text/plain;charset=UTF-8")

@crud.route("/qrcode",methods=["GET"])
def get_qrcode():
    data = request.args.get("data", "")
    return send_file(qrcode(data, mode="raw"), mimetype="image/png")


@crud.route("/DownloadXLS",methods=["GET"])
def get_DownloadXLS():
    f2A=["-","itemcatno","-","-","-","sess","vouchernum","regSDate","invoicenum","name","model","sn","supplier","quantity","price","p_amount","place","-","fund_amount","fund_name","keeper","amount","depr_year","warr_period","note1","note2"]
    datestr= datetime.today().strftime( '%Y-%m-%d')
    wb = load_workbook(filename = 'doc/asset_data_templ.xlsx')
    cate_row_idx={}
    for sn_ in wb.sheetnames:
        cate_row_idx[sn_]=7
    books, next_page_token = get_assest_model().Itemlist(limit=6000)
    for r_ in books:
        sn_=str(r_["cate"])
        
        if sn_ in wb.sheetnames:
            ridx=cate_row_idx[sn_]
            for i, f_ in enumerate(f2A):
                if f_ == "itemcatno":
                    wb[sn_][f"{chr(65+i)}{ridx}"]=str(r_[f_]).zfill(14)
                elif f_ == "regSDate":
                    wb[sn_][f"{chr(65+i)}{ridx}"]=r_[f_].strftime( '%Y-%m-%d')
                elif f_ != "-":
                    wb[sn_][f"{chr(65+i)}{ridx}"]=r_[f_]
                if chr(65+i) in ['O','P','R'] :
                    wb[sn_][f"{chr(65+i)}{ridx}"].number_format = '#,##0.00'    
            wb[sn_].row_dimensions[ridx].height = 30
            cate_row_idx[sn_]=ridx+1
    file = io.BytesIO()
    wb.save(file)
    file.seek(0)
    return send_file(file, attachment_filename=f"asset_{datestr}.xlsx", as_attachment=True)

import copy 
import intWeb.esasset.models_lebaldata as lebaldata
import qrcode as lebalQrcode
from docx import Document
from docx.shared import Inches,Cm,Pt
from docx.enum.section import WD_ORIENT #處理文件的直向/橫向
from docx.table import _Cell

@crud.route('/templdocx')
def templdocx_router():
    document = Document("doc/lebal_templ_schsub951.docx")
    template = document.tables[0]
    tbl = template._tbl
    lebals=lebaldata.lebal_data.split("\n")
    lebals_len=len(lebals)//4

    for i in range(lebals_len):
        new_tbl = copy.deepcopy(tbl)
        # Then we do the replacement
        # replaceText(document, '<>', 'New value')
        #paragrah = document.add_paragraph()
        paragraph =document.paragraphs[0]
        # After that, we add the previously copied table
        paragraph._p.addnext(new_tbl)    

    for i, table in enumerate(document.tables):
        if i>=lebals_len:
            break
        print(i)
        qr_txt=lebals[i*4+3]
        lebal_txt=f"{lebals[i*4+0]}\n{lebals[i*4+1]}\n{lebals[i*4+2]}\n{lebals[i*4+3]}"
        print(lebal_txt)
        hdr_cells = table.rows[0].cells
        file = lebalQrcode.make(qr_txt)
        buf = io.BytesIO()
        file.save(buf)
        buf.seek(0)
        p=hdr_cells[0].paragraphs[0]
        run=p.add_run()
        run.add_picture(buf, width=Cm(1.8)) # width=Inches(1.25))
        p=hdr_cells[1].paragraphs[0]
        cell_run=p.add_run(lebal_txt)
        cell_run.font.name = '新細明體'#'微軟正黑體'
        cell_run.font.size = Pt(8)
        #cell_run.bold = True    
    f = io.BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    return send_file(f, as_attachment=True, attachment_filename='report.doc')    
